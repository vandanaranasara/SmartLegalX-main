from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from transformers import pipeline
import fitz  # also install PyMuPDF library for PDF text extraction
import os
import traceback

from werkzeug.utils import secure_filename
from datetime import datetime
from docx import Document

app = Flask(__name__)
CORS(app, expose_headers='Authorization')

ALLOWED_EXTENSIONS = set(['txt', 'pdf', 'png', 'jpg', 'jpeg', 'gif'])
app.secret_key = "secret key"
app.config['UPLOAD_FOLDER'] = "C:/SmartLegalX-main/flaskback/assets"
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024

summarizer = pipeline("summarization", model="t5-base")

# Helper function to check if a file has an allowed extension
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Summarize text using the Hugging Face T5 model
@app.route('/summarize', methods=['POST'])
def summarize_text():
    try:
        article = request.json.get('article')
        if not article:
            return jsonify({'error': 'Empty input'}), 500

        # Perform summarization using the pipeline
        summary = summarizer(article, max_length=130, min_length=30, do_sample=False)

        # Extract the summary text from the result
        summary_text = summary[0]['summary_text']

        # Create a dictionary to store the summary
        summary_dict = {'summary': summary_text}

        # Return the summary as JSON
        return jsonify(summary_dict)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# Upload a file, extract text from a PDF, and summarize it
@app.route('/upload', methods=['POST'])
def upload_file():
    try:
        # Check if the post request has the file part
        if 'file' not in request.files:
            return jsonify({'message': 'No file part in the request'}), 400

        file = request.files['file']

        if file.filename == '':
            return jsonify({'message': 'No file selected for uploading'}), 400

        if file and allowed_file(file.filename):
            try:
                filename = secure_filename(file.filename)
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(file_path)

                # Extract text from the PDF file
                pdf = fitz.open(file_path)
                text = ""
                for page_num in range(pdf.page_count):
                    page = pdf.load_page(page_num)
                    text += page.get_text()

                # text = text.strip()
                # Perform summarization using the pipeline
                summary = summarizer(text, max_length=130, min_length=30, do_sample=False)

                # Extract the summary text from the result
                summary_text = summary[0]['summary_text']

                return jsonify({'message': 'File successfully uploaded and summarized', 'summary': summary_text}),200
            except Exception as e:
                return jsonify({'message': 'Error processing the uploaded file', 'error': str(e)}), 500
        else:
            return jsonify({'message': 'Allowed file types are txt, pdf, png, jpg, jpeg, gif'}), 400
    except Exception as e:
        return jsonify({'error': str(e)}), 500



'''
generate will template using user form input
'''

def format_date(date):
    try:
        # Parse the formatted date and reformat it as needed
        parsed_date = datetime.strptime(date, '%B %d, %Y')
        formatted_date = parsed_date.strftime('%d-name day of %B %Y-number year')
    except ValueError:
        formatted_date = ''  # Handle invalid date format gracefully if needed
    return formatted_date

@app.route('/generate_will', methods=['POST'])
def generate_document():
    try:
        # Get user input from the form (assuming data is sent as JSON)
        user_data = request.json

        # Load the template document
        template_doc = Document('C:/SmartLegalX-main/flaskback/assets/WillDeed.docx')

        formatted_date = format_date(user_data.get('date', ''))

        # Replace placeholders in the template with user data
        for paragraph in template_doc.paragraphs:
            for run in paragraph.runs:
                run.text = run.text.replace('user-name-input', user_data.get('name', ''))
                run.text = run.text.replace('son-of', user_data.get('sonOf', ''))
                run.text = run.text.replace('abc-address-abc-colony', user_data.get('residence', ''))
                run.text = run.text.replace('age-entry', user_data.get('age', ''))
                run.text = run.text.replace('religion-input', user_data.get('religion', ''))
                run.text = run.text.replace('occupation-name', user_data.get('occupation', ''))

                run.text = run.text.replace('Will-receiver-name', user_data.get('willReceiverName', ''))
                run.text = run.text.replace('will-receiver-address', user_data.get('willReceiverAddress', ''))
                run.text = run.text.replace('will-receiver-age', user_data.get('willReceiverAge', ''))
                run.text = run.text.replace('will-receiver-religion', user_data.get('willReceiverReligion', ''))
                run.text = run.text.replace('will-receiver-occupation', user_data.get('willReceiverOccupation', ''))
                run.text = run.text.replace('Will-receiver2-name', user_data.get('willReceiver2Name', ''))
                run.text = run.text.replace('will-receiver2-address', user_data.get('willReceiver2Address', ''))
                run.text = run.text.replace('will-receiver2-age', user_data.get('willReceiver2Age', ''))
                run.text = run.text.replace('will-receiver2-religion', user_data.get('willReceiver2Religion', ''))
                run.text = run.text.replace('will-receiver2-occupation', user_data.get('willReceiver2Occupation', ''))
                run.text = run.text.replace('Will-receiver3-name', user_data.get('willReceiver3Name', ''))
                run.text = run.text.replace('will-receiver3-address', user_data.get('willReceiver3Address', ''))
                run.text = run.text.replace('will-receiver3-age', user_data.get('willReceiver3Age', ''))
                run.text = run.text.replace('will-receiver3-religion', user_data.get('willReceiver3Religion', ''))
                run.text = run.text.replace('will-receiver3-occupation', user_data.get('willReceiver3Occupation', ''))
          
                
                run.text = run.text.replace('family-members', user_data.get('familyMembers', ''))
                run.text = run.text.replace('PropertyInput1', user_data.get('property', ''))

                
                run.text = run.text.replace('User-input-name-wife', user_data.get('userWife', ''))
                
                run.text = run.text.replace('Child1-name', user_data.get('child1Name', ''))
                run.text = run.text.replace('Child2-name', user_data.get('child2Name', ''))
                run.text = run.text.replace('Child3-name', user_data.get('child3Name', ''))
                
                run.text = run.text.replace('day-name day of month-name month of year-number year', formatted_date)

                run.text = run.text.replace('witness1-name', user_data.get('witness1Name', ''))
                run.text = run.text.replace('witness2-name', user_data.get('witness2Name', ''))

        # Save the filled-out document
        filled_doc_path = 'filled_legal_document.docx'
        template_doc.save(filled_doc_path)

        # Send the filled document back to the user for download
        return send_file(filled_doc_path, as_attachment=True)
    except Exception as e:
        return jsonify({'error': str(e)})


@app.route('/anti_ragging', methods=['POST'])
def anti_ragging_document():
    try:
        # Get user input from the form (assuming data is sent as JSON)
        user_data = request.json

        # Load the template document for the anti-ragging form (adjust the path accordingly)
        template_doc = Document('C:/SmartLegalX-main/flaskback/assets/AntiRagging.docx')

        # Define a dictionary of placeholders and their corresponding user data keys
        placeholders = {
        'student-fullname': 'studentFullName',
        'admission-no': 'admissionNo',
        'parent-name': 'parentName',
        'institute-name': 'instituteName',
        'day': 'declarationDay',  # Example: "declarationDay": "5"
        'month': 'declarationMonth',  # Example: "declarationMonth": "September"
        'year': 'declarationYear',  # Example: "declarationYear": "2023"
        'parent-guardian-fullname': 'parentGuardianFullName',
        'student-fullname2': 'studentFullName',
        'admission-no2': 'admissionNo',
        'institute-name2': 'instituteName',
        'parent-guardian-fullname2': 'parentGuardianFullName',
        'day2': 'declarationDay',  # Example: "declarationDay": "5"
        'month2': 'declarationMonth',  # Example: "declarationMonth": "September"
        'year2': 'declarationYear',  # Example: "declarationYear": "2023"

        'parent-address': 'parentAddress',
        'parent-mobi-no': 'parentMobileNo',
        }


        # Replace placeholders in the template with user data
        for paragraph in template_doc.paragraphs:
            for run in paragraph.runs:
                for placeholder, key in placeholders.items():
                    run.text = run.text.replace(placeholder, user_data.get(key, ''))

        # Save the filled-out anti-ragging form document
        filled_doc_path = 'filled_anti_ragging_form.docx'
        template_doc.save(filled_doc_path)

        # Send the filled document back to the user for download
        return send_file(filled_doc_path, as_attachment=True)
    except Exception as e:
        return jsonify({'error': str(e)})



if __name__ == '__main__':
    app.run(debug=True , port=5000, host="0.0.0.0")
